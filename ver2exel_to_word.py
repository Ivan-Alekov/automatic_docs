import openpyxl
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, sys, time
from docxtpl import DocxTemplate
from docx import Document
from ver2exel_to_word_trybi import trybi_add

if not os.path.isdir('Отчет(ы)'):
    os.mkdir('Отчет(ы)')
os.chdir("Отчет(ы)")

path = "../template3.xlsx"  #  файл можно имееновать как хочется, главное чтобы названия совпадали после ../
# wb = xlrd.open_workbook(path)
# sheet = wb.sheet_by_index(0)
wb1 = openpyxl.reader.excel.load_workbook(filename=path, data_only=True)
wb1.active = 0
sheet1 = wb1.active

#  можно указать в таблице, можно в программе
YEAR = 2022  # можно менять
year_develop = sheet1[f'C5'].value


chapter = sheet1[f'B2'].value
name_city = sheet1[f'B1'].value if sheet1[f'B1'].value != None else '-'
customer = sheet1[f'B3'].value
district = sheet1[f'B4'].value
district2 = ["-".join([j.title() for j in i.split('-')]) if "-" in i else i for i in map(str.title, district.split())]
district2 = district2[0] + " " + district2[1] + "," + district2[2] + " " + district2[3]

type_of_covering = ['асфальтобетон', 'железобетон', 'щебень', 'щебень-гравий', 'грунто-щебень', 'грунт']
type_of_wear = ['капитальный ', 'облегченный', 'переходный', 'низший']

def for_loop_table_list(numb=0):
    pass

t = time.time()
k = 10
while True:
    docs = Document("../tech_report_example.docx")
    if sheet1[f'A{k}'].value == "КОНЕЦ":
        break

    count_none = 0
    if sheet1[f'A{k}'].value == None:
        count_name_road += 1
        k += 1
        continue
    else:
        count_name_road = 0
        while True:
            if sheet1[f'A{k + count_none + 1}'].value == None:
                count_none += 1
            else:
                break
    for_name_road = sheet1[f'A{k - count_name_road}'].value
    for_year = sheet1[f'B{k - count_name_road}'].value
    for_length = round(sum(float(sheet1[f'C{k - count_name_road + i}'].value) for i in range(count_none + 1)), 3)

    result_read_exel = [for_name_road, for_year, for_length,
                        sheet1[f'D{k}'].value, sheet1[f'E{k}'].value, sheet1[f'F{k}'].value,
                        sheet1[f'G{k}'].value, sheet1[f'H{k}'].value, sheet1[f'I{k}'].value,
                        sheet1[f'J{k}'].value, sheet1[f'K{k}'].value, sheet1[f'L{k}'].value,
                        sheet1[f'M{k}'].value, sheet1[f'N{k}'].value, sheet1[f'O{k}'].value,
                        sheet1[f'P{k}'].value,
                        sheet1[f'Q{k}'].value, sheet1[f'R{k}'].value, sheet1[f'S{k}'].value,
                        sheet1[f'T{k}'].value, sheet1[f'U{k}'].value, sheet1[f'V{k}'].value,
                        sheet1[f'W{k}'].value, sheet1[f'X{k}'].value, sheet1[f'Y{k}'].value,
                        sheet1[f'Z{k}'].value, sheet1[f'AA{k}'].value, sheet1[f'AB{k}'].value,
                        sheet1[f'AC{k}'].value, sheet1[f'AG{k}'].value]

    go_word = {'name_road': result_read_exel[0], 'year_develop': year_develop,
               'year': int(result_read_exel[1]) if result_read_exel[1] != None else 'данные отсутствуют',
               'lenght_all': f'{result_read_exel[2]:.3f}'.replace('.', ','),
               'squere': 0.0, "district2": district2,
               'year_check': sheet1[f'C5'].value if sheet1[f'C5'].value != None else f'{YEAR}',
               'name_city': name_city, 'chapter': chapter, 'customer': customer, 'district': district,
               'covering_weight': [], 'lenght_to40': '-', 'lenght_to41_44': '-', 'lenght_to45_59': '-',
               'lenght_to60_65': '-', 'lenght_to66_69': '-', 'lenght_to70_74': '-', 'lenght_to75_84': '-',
               'lenght_to85_11': '-', 'lenght_to11_14': '-',
               'weight_to80': '-', 'weight_to80_99': '-', 'weight_to100_119': '-', 'weight_to120_149': '-',
               'weight_to150_274': '-', 'weight_to275': '-',
               'asphalt': 0.0,
               'asphalt_l': 0.0,
               'reinforced_l': 0.0,
               'rubble_l': 0.0,
               'crushed_stone_l': 0.0,
               'ground_crushed_l': 0.0,
               'priming_l': 0.0,
               'asphalt_w': f'{result_read_exel[9]:.3f}'.replace(".", ',') if result_read_exel[9] != None else '',
               'reinforced_w': f'{result_read_exel[10]:.3f}'.replace(".", ',') if result_read_exel[10] != None else '',
               'rubble_w': f'{result_read_exel[11]:.3f}'.replace(".", ',') if result_read_exel[11] != None else '',
               'crushed_stone_w': f'{result_read_exel[12]:.3f}'.replace(".", ',') if result_read_exel[12] != None else '',
               'ground_crushed_w': f'{result_read_exel[13]:.3f}'.replace(".", ',') if result_read_exel[13] != None else '',
               'priming_w': f'{result_read_exel[14]:.3f}'.replace(".", ',') if result_read_exel[14] != None else '',
               'lead_district': f"{sheet1[f'B{2}'].value}" if sheet1[f'B{2}'].value != None else '',
               'admin_country': f"{sheet1[f'B{3}'].value}" if sheet1[f'B{3}'].value != None else '',
               # 'in_table_w': f'{result_read_exel[26]}' if result_read_exel[26] != None else '-',
               # 'in_table_cond': f'{result_read_exel[16]}' if result_read_exel[16] != None else '-',
               # 'in_table_andgl': f'{result_read_exel[27]}' if result_read_exel[27] != None else '-',
               'AG_type_pedestr': f'{result_read_exel[29]}' if result_read_exel[29] != None else '-',
               'AH_pedestr': f'{sheet1[f"AH{k}"].value}' if sheet1[f'AH{k}'].value != None else '-',
               'AI_pedestr': f'{sheet1[f"AI{k}"].value}' if sheet1[f'AI{k}'].value != None else '-'}

    go_word['covering_weight1'] = go_word['lenght_all'].split(',')[0]
    go_word['covering_weight2'] = go_word['lenght_all'].split(',')[1]

    roads_part_adress = []
    roads_part_lenght = []
    roads_part_weight = []
    cathegory_roads = []
    numb_lines = []
    type_road = []
    roads_speed = []
    power_traffic = []
    roads_group = []

    go_word['count_numb'] = ', '.join(str(i + 1) for i in range(count_none + 1))
    go_word['area'] = f"на участках №{', '.join(str(i + 1) for i in range(count_none + 1))} " \
                      f"по состоянию на {go_word['year_check']}" if count_none > 0 else \
        f"на участке №1 по состоянию на {go_word['year_check']}"

    speed_30 = ['проезд', 'местная дорога']

    lenght_to40 = []
    lenght_to41_44 = []
    lenght_to45_59 = []
    lenght_to60_65 = []
    lenght_to66_69 = []
    lenght_to70_74 = []
    lenght_to75_84 = []
    lenght_to85_11 = []
    lenght_to11_14 = []

    weight_to80 = []
    weight_to80_99 = []
    weight_to100_119 = []
    weight_to120_149 = []
    weight_to150_274 = []
    weight_to275 = []
    ans_name_table_351 = []

    #  Цикл для создания ответа пункта: ТЕХНИЧЕСКАЯ ХАРАКТЕРИСТИКА
    for i in range(count_none + 1):
        lenght_all = list(map(lambda x: float(x) if x != None else 0, [sheet1[f'S{k + i}'].value,
                                                                                  sheet1[f'T{k + i}'].value,
                                                                                  sheet1[f'U{k + i}'].value,
                                                                                  sheet1[f'V{k + i}'].value]))
        l_all = [(ind, znach) for ind, znach in enumerate(lenght_all) if znach != 0]
        ans_name_table_351.append(l_all[0][0])
        weight_all = list(map(lambda x: float(x) if x != None else 0, [sheet1[f'J{k + i}'].value,
                                                                       sheet1[f'K{k + i}'].value,
                                                                       sheet1[f'L{k + i}'].value,
                                                                       sheet1[f'M{k + i}'].value,
                                                                       sheet1[f'N{k + i}'].value,
                                                                       sheet1[f'O{k + i}'].value]))
        w_all = [(ind, znach) for ind, znach in enumerate(weight_all) if znach != 0]

        roads_part_adress.append(f"- Участок дороги  №{i + 1} «{sheet1[f'R{k + i}'].value}»;")
        roads_part_lenght.append(f"       Участок №{i + 1}- {f'{l_all[0][1]:.3f}'.replace('.', ',')}км.")
        roads_part_weight.append(f"       Участок №{i + 1} ({type_of_covering[w_all[0][0]]}) - {f'{w_all[0][1]:.3f}'.replace('.', ',')}м.")
        cathegory_roads.append(f"       Участок №{i + 1}- {sheet1[f'W{k + i}'].value};")
        numb_lines.append(f'       Участок №{i + 1}- 2' if w_all[0][1] >= 4.5 else f'       Участок №{i + 1}- 1')
        type_road.append(f"       Участок №{i + 1}- {type_of_covering[w_all[0][0]]}")
        power_traffic.append(f"       Участок №{i + 1}- {sheet1[f'X{k + i}'].value} авт/сут")
        roads_speed.append(f'       Участок №{i + 1}- 30 км/ч' if sheet1[f"W{k + i}"].value.lower() in speed_30 else\
                                 f'       Участок №{i + 1}- 40 км/ч')
        roads_group.append(f"       Участок №{i + 1}- {sheet1[f'AC{k + i}'].value.capitalize()}")

        #  для таблицы 3.3.1. Ширина проезжей части
        if w_all[0][1] < 4:
            lenght_to40.append(l_all[0][1])
        elif 4 <= w_all[0][1] <= 4.4:
            lenght_to41_44.append(l_all[0][1])
        elif 4.5 <= w_all[0][1] <= 5.9:
            lenght_to45_59.append(l_all[0][1])
        elif 6 <= w_all[0][1] <= 6.5:
            lenght_to60_65.append(l_all[0][1])
        elif 6.6 <= w_all[0][1] <= 6.9:
            lenght_to66_69.append(l_all[0][1])
        elif 7 <= w_all[0][1] <= 7.4:
            lenght_to70_74.append(l_all[0][1])
        elif 7.5 <= w_all[0][1] <= 8.4:
            lenght_to75_84.append(l_all[0][1])
        elif 8.5 <= w_all[0][1] <= 11:
            lenght_to85_11.append(l_all[0][1])
        elif 11.1 <= w_all[0][1] <= 14:
            lenght_to11_14.append(l_all[0][1])

        #  для таблицы 3.2. Ширина земляного полотна
        if w_all[0][1] < 8:
            weight_to80.append(l_all[0][1])
        elif 8 <= w_all[0][1] <= 9.9:
            weight_to80_99.append(l_all[0][1])
        elif 10 <= w_all[0][1] <= 11.9:
            weight_to100_119.append(l_all[0][1])
        elif 12 <= w_all[0][1] <= 14.9:
            weight_to120_149.append(l_all[0][1])
        elif 15 <= w_all[0][1] <= 27.4:
            weight_to150_274.append(l_all[0][1])
        elif 27.5 <= w_all[0][1]:
            weight_to275.append(l_all[0][1])

        #  заполнение длин типов покрытия
        if w_all[0][0] == 0 and l_all[0][0] == 0:
            go_word['asphalt'] = float(go_word['asphalt']) + l_all[0][1]
        elif w_all[0][0] == 0 and l_all[0][0] == 1:
            go_word['asphalt_l'] = float(go_word['asphalt_l']) + l_all[0][1]
        elif w_all[0][0] == 1:
            go_word['reinforced_l'] = float(go_word['reinforced_l']) + l_all[0][1]
        elif w_all[0][0] == 2:
            go_word['rubble_l'] = float(go_word['rubble_l']) + l_all[0][1]
        elif w_all[0][0] == 4:
            go_word['ground_crushed_l'] = float(go_word['ground_crushed_l']) + l_all[0][1]
        elif w_all[0][0] == 5:
            go_word['priming_l'] = float(go_word['priming_l']) + l_all[0][1]

    #  для всех типов покрытий приводим к одному виду
    go_word['asphalt'] = f"{go_word['asphalt']:.3f}".replace('.', ",") if go_word['asphalt'] != 0.0 else '0,0'
    go_word['asphalt_l'] = f"{go_word['asphalt_l']:.3f}".replace('.', ",") if go_word['asphalt_l'] != 0.0 else '0,0'
    go_word['reinforced_l'] = f"{go_word['reinforced_l']:.3f}".replace('.', ",") if go_word['reinforced_l'] != 0.0 else '0,0'
    go_word['rubble_l'] = f"{go_word['rubble_l']:.3f}".replace('.', ",") if go_word['rubble_l'] != 0.0 else '0,0'
    go_word['ground_crushed_l'] = f"{go_word['ground_crushed_l']:.3f}".replace('.', ",") if go_word['ground_crushed_l'] != 0.0 else '0,0'
    go_word['priming_l'] = f"{go_word['priming_l']:.3f}".replace('.', ",") if go_word['priming_l'] != 0.0 else '0,0'

    ot_lenght = go_word['lenght_all'].split(",")[0]
    for_length = go_word['lenght_all'].split(",")[1]
    go_word['roads_part_adress'] = "\n".join(i for i in roads_part_adress) if len(roads_part_adress) > 1 else \
        f"- Участок №1- км 0+000 до км {ot_lenght}+{for_length}"

    go_word['roads_part_lenght'] = '\n' + '\n'.join(i for i in roads_part_lenght) if len(roads_part_lenght) > 1 else \
        f'{go_word["lenght_all"]}км'.replace(".", ",")

    go_word['roads_part_weight'] = '\n' + '\n'.join(i for i in roads_part_weight) if len(roads_part_weight) > 1 else \
        f"{f'{w_all[0][1]:.3f}'}м."

    go_word["cathegory_roads"] = '\n' + '\n'.join(i for i in cathegory_roads) if len(cathegory_roads) > 1 else \
        f'{cathegory_roads[0]}'.split()[-1]

    go_word["numb_lines"] = '\n' + '\n'.join(i for i in numb_lines) if len(numb_lines) > 1 else \
        f'{numb_lines[0]};'.split()[-1]

    go_word["type_road"] = '\n' + '\n'.join(i for i in type_road) if len(type_road) > 1 else \
        f'{type_road[0]};'.split()[-1]

    go_word["power_traffic"] = '\n' + '\n'.join(i for i in power_traffic) if len(power_traffic) > 1 else \
        f'{power_traffic[0]};'.split()[-2] + f'{power_traffic[0]};'.split()[-1]

    go_word["roads_speed"] = '\n' + '\n'.join(i for i in roads_speed) if len(roads_speed) > 1 else \
        f'{roads_speed[0]};'.split()[-2] + f'{roads_speed[0]};'.split()[-1]

    go_word["roads_group"] = '\n' + '\n'.join(i for i in roads_group) if len(roads_group) > 1 else \
        f'{roads_group[0]};'.split()[-1].capitalize()

    #  данные для таблиц 3.2 и 3.3.1 (Ширина земляного полотна и Ширина проезжей части)
    go_word['lenght_to40'] = f'{sum(lenght_to40):.3f}'.replace(".", ",") if lenght_to40 else '-'
    go_word['lenght_to41_44'] = f'{sum(lenght_to41_44):.3f}'.replace(".", ",") if lenght_to41_44 else '-'
    go_word['lenght_to45_59'] = f'{sum(lenght_to45_59):.3f}'.replace(".", ",") if lenght_to45_59 else '-'
    go_word['lenght_to60_65'] = f'{sum(lenght_to60_65):.3f}'.replace(".", ",") if lenght_to60_65 else '-'
    go_word['lenght_to66_69'] = f'{sum(lenght_to66_69):.3f}'.replace(".", ",") if lenght_to66_69 else '-'
    go_word['lenght_to70_74'] = f'{sum(lenght_to70_74):.3f}'.replace(".", ",") if lenght_to70_74 else '-'
    go_word['lenght_to75_84'] = f'{sum(lenght_to75_84):.3f}'.replace(".", ",") if lenght_to75_84 else '-'
    go_word['lenght_to85_11'] = f'{sum(lenght_to85_11):.3f}'.replace(".", ",") if lenght_to85_11 else '-'
    go_word['lenght_to11_14'] = f'{sum(lenght_to11_14):.3f}'.replace(".", ",") if lenght_to11_14 else '-'

    go_word['weight_to80'] = f'{sum(weight_to80):.3f}'.replace(".", ",") if weight_to80 else '-'
    go_word['weight_to80_99'] = f'{sum(weight_to80_99):.3f}'.replace(".", ",") if weight_to80_99 else '-'
    go_word['weight_to100_119'] = f'{sum(weight_to100_119):.3f}'.replace(".", ",") if weight_to100_119 else '-'
    go_word['weight_to120_149'] = f'{sum(weight_to120_149):.3f}'.replace(".", ",") if weight_to120_149 else '-'
    go_word['weight_to150_274'] = f'{sum(weight_to150_274):.3f}'.replace(".", ",") if weight_to150_274 else '-'
    go_word['weight_to275'] = f'{sum(weight_to275):.3f}'.replace(".", ",") if weight_to275 else '-'

    #print(go_word["asphalt"] ,go_word["reinforced_l"], go_word["asphalt_l"], go_word["rubble_l"], go_word["ground_crushed_l"], go_word["priming_l"])

    #  инофрмация для таблицы 3.3.2 Протяженность покрытий
    go_word["best_road"] = float(go_word["asphalt"].replace(",", ".")) + float(go_word["reinforced_l"].replace(",", "."))
    go_word["best_road"] = f'{go_word["best_road"]:.3f}'.replace(".", ",") if go_word["best_road"] != 0 else "-"

    go_word["middle_road"] = go_word["rubble_l"] if go_word["rubble_l"] != '0,0' else "-"

    go_word["bad_road"] = float(go_word["ground_crushed_l"].replace(",", ".")) + float(go_word["priming_l"].replace(",", "."))
    go_word["bad_road"] = f'{go_word["bad_road"]:.3f}'.replace(".", ",") if go_word["bad_road"] != 0 else "-"

    go_word["lite_road"] = go_word["asphalt_l"] if go_word["asphalt_l"] != '0,0' else "-"

    go_word["asphalt"] = go_word["asphalt"] if go_word["asphalt"] != '0,0' else ""
    go_word["reinforced_l"] = go_word["reinforced_l"] if go_word["reinforced_l"] != '0,0' else ""
    go_word["asphalt_l"] = go_word["asphalt_l"] if go_word["asphalt_l"] != '0,0' else ""
    go_word["rubble_l"] = go_word["rubble_l"] if go_word["rubble_l"] != '0,0' else ""
    go_word["ground_crushed_l"] = go_word["ground_crushed_l"] if go_word["ground_crushed_l"] != '0,0' else ""
    go_word["priming_l"] = go_word["priming_l"] if go_word["priming_l"] != '0,0' else ""

    # #  данные для таблицы 3.5.1 (Сведения о транспортно-эксплуатационном состоянии проезжей части)
    # ans_from_table_3_5_1_rubble = sum(1 for i in range(count_none + 1) if sheet1[f'F{k + i}'].value != None)
    # ans_from_table_3_5_1_asf = sum(1 for i in range(count_none + 1) if sheet1[f'D{k + i}'].value != None or sheet1[f'E{k + i}'].value != None)

    for ind, numb in enumerate(docs.tables):
        if ind == 1:
            #  таблица ОБЩИЕ ДАННЫЕ
            for i in range(count_none + 1):
                lenght_all = list(map(lambda x: float(x) if x != None else 0, [sheet1[f'S{k + i}'].value,
                                                                               sheet1[f'T{k + i}'].value,
                                                                               sheet1[f'U{k + i}'].value,
                                                                               sheet1[f'V{k + i}'].value]))
                l_all = [(ind, znach) for ind, znach in enumerate(lenght_all) if znach != 0]

                style = docs.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                # docs.tables[ind].cell(3, 2).text = f"0+{go_word['lenght_all'].split(',')[1]}" if\
                #     float(go_word['lenght_all'].split(",")[0]) <= 0 else\
                #     f"{go_word['lenght_all'].split(',')[0]}+{go_word['lenght_all'].split(',')[1]}"
                # docs.tables[ind].cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                if i != count_none:
                    docs.tables[ind].add_row().cells
                docs.tables[ind].cell(i + 3, 0).text = f"0,000"
                docs.tables[ind].cell(i + 3, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 1).text = f"{l_all[0][1]:.3f}".replace(".", ",")
                docs.tables[ind].cell(i + 3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 2).text = f"{l_all[0][1]:.3f}".replace(".", ",")
                docs.tables[ind].cell(i + 3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 3).text = f"-"
                docs.tables[ind].cell(i + 3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 4).text = f"{l_all[0][1]:.3f}".replace(".", ",")
                docs.tables[ind].cell(i + 3, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 5).text = f"{l_all[0][1]:.3f}".replace(".", ",")
                docs.tables[ind].cell(i + 3, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 6).text = f"{l_all[0][1]:.3f}".replace(".", ",")
                docs.tables[ind].cell(i + 3, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 7).text = f"-"
                docs.tables[ind].cell(i + 3, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


        if ind == 3:
            # таблица Категория дороги (участка), подъездов
            for i in range(count_none + 1):
                lenght_all = list(map(lambda x: float(x) if x != None else 0, [sheet1[f'S{k + i}'].value,
                                                                               sheet1[f'T{k + i}'].value,
                                                                               sheet1[f'U{k + i}'].value,
                                                                               sheet1[f'V{k + i}'].value]))
                l_all = [(ind, znach) for ind, znach in enumerate(lenght_all) if znach != 0]
                style = docs.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                if count_none == 0:
                    docs.tables[ind].cell(2, 2).text = f"0+{go_word['lenght_all'].split(',')[1]}" if\
                    float(go_word['lenght_all'].split(",")[0]) < 0 else\
                    f"{go_word['lenght_all'].split(',')[0]}+{go_word['lenght_all'].split(',')[1]}"
                    docs.tables[ind].cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    docs.tables[ind].cell(2, 3).text = f"{sheet1[f'W{k + i}'].value}"
                    docs.tables[ind].cell(2, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    break

                docs.tables[ind].cell(2, 2).text = f"0+{go_word['lenght_all'].split(',')[1]}" if\
                    float(go_word['lenght_all'].split(",")[0]) < 0 else\
                    f"{go_word['lenght_all'].split(',')[0]}+{go_word['lenght_all'].split(',')[1]}"
                docs.tables[ind].cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].add_row().cells
                docs.tables[ind].cell(i + 3, 0).text = f"- Участок дороги  №{i + 1} «{sheet1[f'R{k + i}'].value}»"
                docs.tables[ind].cell(i + 3, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 1).text = f"0+000"
                docs.tables[ind].cell(i + 3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                l_all1 = f'{l_all[0][1]:.3f}'.split('.')
                docs.tables[ind].cell(i + 3, 2).text = f"{l_all1[0]}+{l_all1[1]}"
                docs.tables[ind].cell(i + 3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 3, 3).text = f"{sheet1[f'W{k + i}'].value}"
                docs.tables[ind].cell(i + 3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        if ind == 4:
            #  таблица Первичные дорожные организации
            for i in range(count_none + 1):
                lenght_all = list(map(lambda x: float(x) if x != None else 0, [sheet1[f'S{k + i}'].value,
                                                                               sheet1[f'T{k + i}'].value,
                                                                               sheet1[f'U{k + i}'].value,
                                                                               sheet1[f'V{k + i}'].value]))
                l_all = [(ind, znach) for ind, znach in enumerate(lenght_all) if znach != 0]
                style = docs.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                if count_none == 0:
                    docs.tables[ind].cell(3, 5).text = f"0+{go_word['lenght_all'].split(',')[1]}" if \
                        float(go_word['lenght_all'].split(",")[0]) <= 0 else \
                        f"{go_word['lenght_all'].split(',')[0]}+{go_word['lenght_all'].split(',')[1]}"
                    docs.tables[ind].cell(3, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    break

                docs.tables[ind].cell(3, 5).text = f"0+{go_word['lenght_all'].split(',')[1]}" if\
                    float(go_word['lenght_all'].split(",")[0]) <= 0 else\
                    f"{go_word['lenght_all'].split(',')[0]}+{go_word['lenght_all'].split(',')[1]}"
                docs.tables[ind].cell(3, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].add_row().cells
                docs.tables[ind].cell(i + 4, 0).text = f"{YEAR}"
                docs.tables[ind].cell(i + 4, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 1).text = f"{go_word['admin_country']}"
                docs.tables[ind].cell(i + 4, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 2).text = f"-"
                docs.tables[ind].cell(i + 4, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 3).text = f"-"
                docs.tables[ind].cell(i + 4, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 4).text = f"0+000"
                docs.tables[ind].cell(i + 4, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                l_all1 = f'{l_all[0][1]:.3f}'.split('.')
                docs.tables[ind].cell(i + 4, 5).text = f"{l_all1[0]}+{l_all1[1]}"
                docs.tables[ind].cell(i + 4, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 6).text = f"{l_all[0][1]:.3f}".replace(".", ",")
                docs.tables[ind].cell(i + 4, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 7).text = f"-"
                docs.tables[ind].cell(i + 4, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 8).text = f"-"
                docs.tables[ind].cell(i + 4, 8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].cell(i + 4, 9).text = f"-"
                docs.tables[ind].cell(i + 4, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        if ans_name_table_351 and ind == 11:
            #  3.5.1 Сведения о транспортно-эксплуатационном состоянии проезжей части
            style = docs.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.bold = None
            font.size = Pt(12)
            prom = ['Ширина полос движения', 'Состояние покрытия проезжей части', 'Продольный уклон',
                    'Продольная ровность', 'Поперечная ровность (коллейность)', 'Коэффициент сцепления']
            count_line_m = 0
            for line in range(count_none + 1):
                count_line = line * 6 + 2 + count_line_m
                if count_none != 0:
                    docs.tables[ind].add_row().cells
                    docs.tables[ind].cell(count_line, 0).text = f'{roads_part_adress[line][2:]}'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    count_line += 1

                docs.tables[ind].add_row().cells
                docs.tables[ind].cell(count_line, 0).text = f'1'
                docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 1).text = f'{prom[0]}'
                docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 2).text = f'{sheet1[f"AA{k + line}"].value}'
                docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 3).text = f'Требованиям СП 42.13330.2016 с учетом требований п.5.2.2.1 ОДМ 218.4.039-2018'
                docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                if count_none != 0:
                    docs.tables[ind].cell(count_line - 1, 0).merge(docs.tables[ind].cell(count_line - 1, 3))

                docs.tables[ind].add_row().cells
                count_line += 1
                docs.tables[ind].cell(count_line, 0).text = f'2'
                docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 1).text = f'{prom[1]}'
                docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 2).text = f'{sheet1[f"Q{k + line}"].value}'
                docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 3).text = f'Требованиям ОДМ 218.4.039-2018'
                docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                docs.tables[ind].add_row().cells
                count_line += 1
                docs.tables[ind].cell(count_line, 0).text = f'3'
                docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 1).text = f'{prom[2]}'
                docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 2).text = f'{sheet1[f"AB{k + line}"].value}'
                docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                docs.tables[ind].cell(count_line, 3).text = f'Требованиям СП 42.13330.2016 с учетом требований п.5.2.2.5 ОДМ 218.4.039-2018'
                docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                #  заполнение для переходного или низшего
                if ans_name_table_351[line] == 2 or ans_name_table_351[line] == 3:
                    docs.tables[ind].add_row().cells
                    count_line += 1
                    docs.tables[ind].cell(count_line, 0).text = f'4'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 1).text = f"{prom[3]}"
                    docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 2).text = f"не определяется"
                    docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 3).text = f"Требованиям СП 42.13330.2016 с учетом требований п.5.2.2.5 ОДМ 218.4.039-2018"
                    docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    docs.tables[ind].add_row().cells
                    count_line += 1
                    docs.tables[ind].cell(count_line, 0).text = f'5'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 1).text = f"{prom[4]}"
                    docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 2).text = f"не определяется"
                    docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 3).text = f"Требованиям СП 42.13330.2016 с учетом требований п.5.2.2.5 ОДМ 218.4.039-2018"
                    docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    docs.tables[ind].add_row().cells
                    count_line += 1
                    docs.tables[ind].cell(count_line, 0).text = f'6'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 1).text = f"{prom[5]}"
                    docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 2).text = f"не определяется"
                    docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 3).text = f"Требованиям СП 42.13330.2016 с учетом требований п.5.2.2.5 ОДМ 218.4.039-2018"
                    docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif ans_name_table_351[line] == 0 or ans_name_table_351[line] == 1:
                    docs.tables[ind].add_row().cells
                    count_line += 1
                    docs.tables[ind].cell(count_line, 0).text = f'4'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 1).text = f"{prom[3]}"
                    docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 2).text = f'{sheet1[f"AD{k + line}"].value}' if sheet1[f"AD{k + line}"].value is not None else 'не соотв.'
                    docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 3).text = f"Требованиям ГОСТ 50597-2017"
                    docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    docs.tables[ind].add_row().cells
                    count_line += 1
                    docs.tables[ind].cell(count_line, 0).text = f'5'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 1).text = f"{prom[4]}"
                    docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 2).text = f'{sheet1[f"AE{k + line}"].value}' if sheet1[f"AE{k + line}"].value is not None else 'не соотв.'
                    docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 3).text = f"Требованиям ГОСТ 50597-2017"
                    docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    docs.tables[ind].add_row().cells
                    count_line += 1
                    docs.tables[ind].cell(count_line, 0).text = f'6'
                    docs.tables[ind].cell(count_line, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 1).text = f"{prom[5]}"
                    docs.tables[ind].cell(count_line, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 2).text = f'{sheet1[f"AF{k + line}"].value}' if sheet1[f"AF{k + line}"].value is not None else 'не соотв.'
                    docs.tables[ind].cell(count_line, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    docs.tables[ind].cell(count_line, 3).text = f"Требованиям ГОСТ 50597-2017"
                    docs.tables[ind].cell(count_line, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                count_line_m += 1
    for i in range(count_none + 1):
        go_word['squere'] = float(go_word['squere']) + sheet1[f"P{k + i}"].value
    go_word['squere'] = f'{go_word["squere"]:.3f}'.replace('.', ',')

    if not os.path.isdir(go_word['name_road']):
        os.mkdir(go_word['name_road'])
    os.chdir(go_word['name_road'])
    docs.save(f'тех.отчет {go_word["name_road"]}.docx')

    doc = DocxTemplate(f'тех.отчет {go_word["name_road"]}.docx')
    doc.render(go_word)
    doc.save(f'тех.отчет {go_word["name_road"]}.docx')

    os.chdir("..")

    k += 1
    print(k - 10, result_read_exel[0])

wb1.close()
try:
    trybi_add()
except:
    print("Нет файла трубы или что-то опшло не так))))")


print((time.time() - t) / 60, "=минут")