import openpyxl
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, sys, time
from docxtpl import DocxTemplate
from docx import Document

def trybi_add():
    path = "../Трубы.xlsx"
    wb1 = openpyxl.reader.excel.load_workbook(filename=path, data_only=True)
    wb1.active = 0
    sheet1 = wb1.active
    t = time.time()
    k = 3

    while True:
        if sheet1[f'A{k}'].value == "КОНЕЦ":
            break

        count_none = 0
        if sheet1[f'A{k}'].value == None:
            count_name_road += 1
            k += 1
            continue
        else:
            count_name_road = 0
            while True:
                if sheet1[f'A{k + count_none + 1}'].value == None:
                    count_none += 1
                else:
                    break
        for_name_road = sheet1[f'A{k - count_name_road}'].value.split()
        for_name_road = for_name_road[0].lower() + " " + for_name_road[1]
        if not os.path.isdir(for_name_road):
            k += count_none + 1
            continue
        os.chdir(for_name_road)
        docs = Document(f'тех.отчет {for_name_road}.docx')
        ind = 36
        count_metal = 0
        count_metalcru = 0
        style = docs.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(8)
        for i in range(1, count_none + 2):
            if i != count_none + 2 - 1:
                docs.tables[ind].add_row().cells
            docs.tables[ind].cell(i, 1).text = f"{sheet1[f'B{k + i - 1}'].value}" if sheet1[f'B{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 1).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            docs.tables[ind].cell(i, 2).text = f"{sheet1[f'C{k + i - 1}'].value}" if sheet1[f'C{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 2).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            docs.tables[ind].cell(i, 3).text = f"{sheet1[f'D{k + i - 1}'].value}" if sheet1[f'D{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 3).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            docs.tables[ind].cell(i, 4).text = f"{sheet1[f'E{k + i - 1}'].value}" if sheet1[f'E{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 4).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            docs.tables[ind].cell(i, 5).text = f"{sheet1[f'F{k + i - 1}'].value}" if sheet1[f'F{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 5).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            docs.tables[ind].cell(i, 6).text = f"{sheet1[f'G{k + i - 1}'].value}" if sheet1[f'G{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 6).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            docs.tables[ind].cell(i, 7).text = f"{sheet1[f'H{k + i - 1}'].value}" if sheet1[f'H{k + i - 1}'].value != None else "-"
            docs.tables[ind].cell(i, 7).paragraphs[
                0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            if sheet1[f'D{k + i - 1}'].value == 'металлическая':
                count_metal += 1
            elif sheet1[f'D{k + i - 1}'].value == 'железобетонная':
                count_metalcru += 1
        all_metal = count_metal + count_metalcru

        docs.tables[ind].cell(1, 0).merge(docs.tables[ind].cell(1 + count_none, 0))
        docs.tables[ind].cell(1, 0).text = f"{for_name_road}"
        docs.tables[ind].cell(1, 0).paragraphs[
            0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        docs.tables[ind].style = 'Table Grid'

        ind2 = 24
        docs.tables[ind2].cell(2, 1).text = str(count_metal) if count_metal else "-"
        docs.tables[ind2].cell(2, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        docs.tables[ind2].cell(2, 2).text = str(count_metal) if count_metal else "-"
        docs.tables[ind2].cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        docs.tables[ind2].cell(4, 1).text = str(count_metalcru) if count_metalcru else "-"
        docs.tables[ind2].cell(4, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        docs.tables[ind2].cell(4, 2).text = str(count_metalcru) if count_metalcru else "-"
        docs.tables[ind2].cell(4, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        docs.tables[ind2].cell(14, 1).text = str(all_metal) if all_metal else "-"
        docs.tables[ind2].cell(14, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        docs.tables[ind2].cell(14, 2).text = str(all_metal) if all_metal else "-"
        docs.tables[ind2].cell(14, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        k += count_none + 1


        docs.save(f'тех.отчет {for_name_road}.docx')

        os.chdir("..")
    print((time.time() - t) / 60, "=минут")


